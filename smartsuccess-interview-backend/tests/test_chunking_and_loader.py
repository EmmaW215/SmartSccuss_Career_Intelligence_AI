"""
Phase 3 Track A-1 unit tests — chunking policy + document loader (PRD 03 §A4).
"""

import pytest

from app.rag.chunking import CHUNK_POLICIES, chunk_document, content_hash
from app.rag.document_loader import detect_doc_type, extract_text, load_documents

RESUME_FIXTURE = "\n\n".join(
    [
        f"Section {i}: Led the development of distributed systems handling "
        "millions of requests per day. Designed APIs, mentored engineers, "
        "and improved P95 latency by 40% through caching and connection "
        "pooling. Technologies: Python, FastAPI, PostgreSQL, Redis, Docker."
        for i in range(30)
    ]
)

JD_FIXTURE = "\n".join(
    [
        f"- Requirement {i}: 5+ years of backend engineering experience with "
        "Python and cloud infrastructure; responsibilities include system "
        "design, code review, and production operations."
        for i in range(40)
    ]
)


def _token_count(text: str) -> int:
    import tiktoken

    return len(tiktoken.get_encoding("cl100k_base").encode(text))


class TestChunkingPolicy:
    def test_resume_chunks_respect_size_limit(self):
        chunks = chunk_document(
            RESUME_FIXTURE, doc_type="resume", user_id="u1", filename="resume.pdf"
        )
        assert len(chunks) > 1, "long resume must be split"
        limit = CHUNK_POLICIES["resume"].chunk_size
        for chunk in chunks:
            assert _token_count(chunk["content"]) <= limit

    def test_jd_uses_denser_policy(self):
        chunks = chunk_document(JD_FIXTURE, doc_type="job_description", user_id="u1")
        limit = CHUNK_POLICIES["job_description"].chunk_size
        assert len(chunks) > 1
        for chunk in chunks:
            assert _token_count(chunk["content"]) <= limit

    def test_chunks_have_overlap(self):
        chunks = chunk_document(RESUME_FIXTURE, doc_type="resume", user_id="u1")
        # Consecutive chunks share text (overlap policy is 64 tokens)
        first, second = chunks[0]["content"], chunks[1]["content"]
        tail = first[-80:]
        assert any(part in second for part in (tail, first[-40:], first[-20:]))

    def test_metadata_complete_on_every_chunk(self):
        chunks = chunk_document(
            RESUME_FIXTURE,
            doc_type="resume",
            user_id="user_42",
            filename="emma_resume.pdf",
        )
        required = {"user_id", "doc_id", "doc_type", "filename", "chunk_index", "created_at"}
        for i, chunk in enumerate(chunks):
            md = chunk["metadata"]
            assert required <= set(md.keys())
            assert md["user_id"] == "user_42"
            assert md["doc_type"] == "resume"
            assert md["filename"] == "emma_resume.pdf"
            assert md["chunk_index"] == i
            assert md["created_at"]

    def test_question_is_atomic(self):
        question = "Tell me about a time you had to debug a production incident. " * 50
        chunks = chunk_document(question, doc_type="question", user_id="")
        assert len(chunks) == 1, "questions are never split"

    def test_empty_text_returns_no_chunks(self):
        assert chunk_document("", doc_type="resume", user_id="u1") == []
        assert chunk_document("   ", doc_type="resume", user_id="u1") == []

    def test_content_hash_is_stable_and_whitespace_insensitive(self):
        assert content_hash("hello world") == content_hash("  hello world  ")
        assert content_hash("hello world") != content_hash("hello mars")


class TestDocumentLoader:
    def test_plain_text_extraction(self):
        text = extract_text("Just plain text résumé".encode("utf-8"), "notes.txt")
        assert "plain text" in text

    def test_docx_extraction_roundtrip(self):
        import io

        import docx

        document = docx.Document()
        document.add_paragraph("Senior ML Engineer with experience in RAG systems.")
        document.add_paragraph("Education: M.Sc. Computer Science.")
        buffer = io.BytesIO()
        document.save(buffer)

        text = extract_text(buffer.getvalue(), "my_resume.docx")
        assert "Senior ML Engineer" in text
        assert "Education" in text

    def test_detect_doc_type_by_filename(self):
        assert detect_doc_type("emma_resume.pdf", "") == "resume"
        assert detect_doc_type("backend_jd.txt", "") == "job_description"

    def test_detect_doc_type_by_content(self):
        assert detect_doc_type("file1.txt", "experience education skills employment") == "resume"
        assert detect_doc_type("file2.txt", "requirements responsibilities qualifications") == "job_description"
        assert detect_doc_type("file3.txt", "random meeting notes") == "supporting"

    def test_load_documents_skips_empty_and_normalizes(self):
        docs = load_documents(
            [
                {"filename": "resume.txt", "content": b"skills experience education employment record"},
                {"filename": "empty.txt", "content": b""},
                {"filename": "inline.txt", "content": "already a string with requirements responsibilities qualifications"},
            ]
        )
        assert len(docs) == 2
        assert docs[0]["doc_type"] == "resume"
        assert docs[1]["doc_type"] == "job_description"
        assert all(set(d.keys()) == {"filename", "text", "doc_type"} for d in docs)
